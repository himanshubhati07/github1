import os
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

with open("ai_changes.md") as f:
    ai_text = f.read()

# API test results come from the xlsx report written in the previous step — the
# single source of truth for final endpoint status (no intermediate log files).
logs_text = ""
if os.path.isfile("api_test_report.xlsx"):
    import openpyxl
    wb_in = openpyxl.load_workbook("api_test_report.xlsx", data_only=True)
    ws_in = wb_in.active
    rows = [r for r in ws_in.iter_rows(values_only=True) if any(c is not None for c in r)]
    if rows:
        header, *data = rows
        lines = ["\t".join(str(c) if c is not None else "" for c in header)]
        for r in data:
            lines.append("\t".join(str(c) if c is not None else "" for c in r))
        logs_text = "\n".join(lines)

sections = dict()
cur_section = None
cur_lines = []
for line in ai_text.splitlines():
    if line.startswith("COMMIT_MESSAGE:"):
        sections["commit"] = line.replace("COMMIT_MESSAGE:", "").strip()
    elif line.startswith("## "):
        if cur_section is not None and cur_lines:
            sections[cur_section] = "\n".join(cur_lines).strip()
        cur_section = line[3:].strip()
        cur_lines = []
    elif cur_section is not None:
        cur_lines.append(line)
if cur_section is not None and cur_lines:
    sections[cur_section] = "\n".join(cur_lines).strip()

def add_h1(doc, text):
    p = doc.add_heading(text, 1)
    for run in p.runs:
        run.font.color.rgb = RGBColor(0x2F, 0x54, 0x96)

def add_bullet(doc, text):
    doc.add_paragraph(text, style="List Bullet")

doc = Document()
title = doc.add_heading("AI Changes Report", 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
for run in title.runs:
    run.font.color.rgb = RGBColor(0x2F, 0x54, 0x96)

add_h1(doc, "Summary")
doc.add_paragraph(sections.get("commit", "(no commit message)"))

for sec in ["Features Added", "Files Modified", "Files Added",
            "Secrets Extracted", "Secrets Moved", "DB URLs Resolved",
            "Test Results Summary"]:
    content = sections.get(sec, "").strip()
    if not content:
        continue
    add_h1(doc, sec)
    for line in content.splitlines():
        line = line.strip()
        if not line:
            continue
        if line[:2] in ("- ", "* ", "+ "):
            add_bullet(doc, line[2:])
        else:
            doc.add_paragraph(line)

add_h1(doc, "Logical Changes Explanation")
features = sections.get("Features Added", "").strip()
modified  = sections.get("Files Modified", "").strip()
added     = sections.get("Files Added", "").strip()
if features:
    doc.add_paragraph("New capabilities introduced into the existing project:")
    for line in features.splitlines():
        line = line.strip().lstrip("-*+ ")
        if line:
            add_bullet(doc, line)
if modified:
    doc.add_paragraph("Existing files extended or updated to support the new features:")
    for line in modified.splitlines():
        line = line.strip().lstrip("-*+ ")
        if line:
            add_bullet(doc, line)
if added:
    doc.add_paragraph("New files created as part of this change:")
    for line in added.splitlines():
        line = line.strip().lstrip("-*+ ")
        if line:
            add_bullet(doc, line)

if logs_text.strip():
    add_h1(doc, "Test Run Logs")
    para = doc.add_paragraph()
    run  = para.add_run(logs_text.strip())
    run.font.name = "Courier New"
    run.font.size = Pt(8)

doc.save("changes_report.docx")
print("Saved: changes_report.docx")
